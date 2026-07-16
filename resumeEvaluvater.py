import json
import os
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from docx import Document
from groq import Groq
from pydantic import BaseModel, Field
from pypdf import PdfReader


BASE_DIR = Path(__file__).resolve().parent
RESUME_FOLDER = BASE_DIR / "resumes"
OUTPUT_FOLDER = BASE_DIR / "reports"
MODEL = "openai/gpt-oss-120b"
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

JOB_DESCRIPTION = """
Description
Amazon is hiring Software Development Engineers (SDE-I) to build scalable
customer-facing products and services. Candidates should be comfortable with
modern programming languages, data structures, algorithms, object-oriented
design, cloud-native systems, microservices, CI/CD, monitoring, debugging,
and collaborating in an agile environment.

Basic Qualifications
- Experience with Java, Python, C++, C#, Go, Rust, or TypeScript.
- Data structures, algorithms, and object-oriented design knowledge.
- A completed or in-progress bachelor's degree in Computer Science or a
  related STEM field.

Preferred Qualifications
- Technical internship or relevant project experience.
- AWS/cloud, SQL/NoSQL databases, Git, AI developer tools, debugging, or
  troubleshooting experience.
- Strong problem solving, written communication, and ability to learn.
"""


class JobDescription(BaseModel):
    role: str
    required_skills: list[str] = Field(default_factory=list)
    preferred_skills: list[str] = Field(default_factory=list)
    minimum_experience_years: float | None = None
    education_requirements: list[str] = Field(default_factory=list)
    responsibilities: list[str] = Field(default_factory=list)


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = Field(default_factory=list)
    experiences: list[Experience] = Field(default_factory=list)
    education: list[str] = Field(default_factory=list)
    projects: list[str] = Field(default_factory=list)
    certifications: list[str] = Field(default_factory=list)


class MatchResult(BaseModel):
    overall_score: float = Field(ge=0, le=100)
    skill_score: float = Field(ge=0, le=100)
    experience_score: float = Field(ge=0, le=100)
    education_score: float = Field(ge=0, le=100)
    matching_skills: list[str] = Field(default_factory=list)
    missing_required_skills: list[str] = Field(default_factory=list)
    matching_preferred_skills: list[str] = Field(default_factory=list)
    experience_requirement_met: bool | None = None
    strengths: list[str] = Field(default_factory=list)
    concerns: list[str] = Field(default_factory=list)
    interview_recommendation: str
    verdict: str


class CandidateReport(BaseModel):
    rank: int | None = None
    file_name: str
    candidate: Resume
    evaluation: MatchResult


def get_client() -> Groq:
    load_dotenv(BASE_DIR / ".env")
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY is missing. Add it to a .env file or your environment.")
    return Groq(api_key=api_key)


def request_json(client: Groq, messages: list[dict[str, str]], schema: dict[str, Any]) -> dict[str, Any]:
    """Call the model with one retry for temporary API or JSON errors."""
    last_error: Exception | None = None
    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                response_format={"type": "json_object"},
            )
            content = response.choices[0].message.content
            if not content:
                raise ValueError("The model returned an empty response.")
            return json.loads(content)
        except Exception as error:
            last_error = error
            if attempt == 0:
                print(f"  Retrying after API/JSON error: {error}")
                time.sleep(2)
    raise RuntimeError(f"Model request failed: {last_error}")


def extract_job(client: Groq, description: str) -> JobDescription:
    schema = JobDescription.model_json_schema()
    messages = [{
        "role": "system",
        "content": f"""Extract a job description into JSON matching this schema:\n{json.dumps(schema)}\n
Return only JSON. Do not invent requirements. Use empty lists or null when unavailable.""",
    }, {
        "role": "user",
        "content": description,
    }]
    return JobDescription(**request_json(client, messages, schema))


def parse_resume(client: Groq, resume_text: str) -> Resume:
    schema = Resume.model_json_schema()
    messages = [{
        "role": "system",
        "content": f"""You are a careful resume parser. Return only JSON matching this schema:\n{json.dumps(schema)}\n
Extract skills from the entire resume, including projects and internships. Never invent information. Use null or empty lists when data is absent.""",
    }, {
        "role": "user",
        "content": resume_text,
    }]
    return Resume(**request_json(client, messages, schema))


def evaluate_candidate(client: Groq, job: JobDescription, resume: Resume) -> MatchResult:
    schema = MatchResult.model_json_schema()
    messages = [{
        "role": "system",
        "content": f"""You are a fair technical recruiter. Evaluate the candidate strictly using the supplied job and resume.
Return only JSON matching this schema:\n{json.dumps(schema)}\n
Scoring: required skills 45%, relevant experience/projects 30%, education 10%, preferred skills 10%, communication/other evidence 5%. Do not penalize candidates for protected characteristics. Explain evidence-based strengths and concerns concisely.""",
    }, {
        "role": "user",
        "content": f"JOB:\n{job.model_dump_json(indent=2)}\n\nRESUME:\n{resume.model_dump_json(indent=2)}",
    }]
    return MatchResult(**request_json(client, messages, schema))


def read_pdf(file_path: Path) -> str:
    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def read_docx(file_path: Path) -> str:
    document = Document(file_path)
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
    return "\n".join(parts)


def read_resume(file_path: Path) -> str:
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    if file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    raise ValueError(f"Unsupported file type: {file_path.suffix}")


def print_candidate(candidate: CandidateReport) -> None:
    result = candidate.evaluation
    print(f"\n#{candidate.rank}  {candidate.candidate.name or 'Unknown candidate'}  |  {result.overall_score:.1f}/100")
    print(f"File: {candidate.file_name}")
    print(f"Scores: skills {result.skill_score:.0f} | experience {result.experience_score:.0f} | education {result.education_score:.0f}")
    print(f"Recommendation: {result.interview_recommendation}")
    print(f"Matching skills: {', '.join(result.matching_skills) or 'None identified'}")
    print(f"Missing skills: {', '.join(result.missing_required_skills) or 'None identified'}")
    print(f"Strengths: {'; '.join(result.strengths) or 'None identified'}")
    print(f"Concerns: {'; '.join(result.concerns) or 'None identified'}")
    print(f"Verdict: {result.verdict}")


def save_report(job: JobDescription, candidates: list[CandidateReport], failures: list[dict[str, str]]) -> Path:
    OUTPUT_FOLDER.mkdir(exist_ok=True)
    report_path = OUTPUT_FOLDER / f"resume_evaluation_{datetime.now():%Y%m%d_%H%M%S}.json"
    report = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "job": job.model_dump(),
        "summary": {
            "processed": len(candidates),
            "failed": len(failures),
            "top_candidate": candidates[0].candidate.name if candidates else None,
        },
        "candidates": [candidate.model_dump() for candidate in candidates],
        "failures": failures,
    }
    report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    return report_path


def main() -> None:
    if not RESUME_FOLDER.is_dir():
        raise FileNotFoundError(f"Resume folder not found: {RESUME_FOLDER}")

    files = sorted(path for path in RESUME_FOLDER.iterdir() if path.suffix.lower() in SUPPORTED_EXTENSIONS)
    if not files:
        raise FileNotFoundError(f"No PDF or DOCX resumes found in: {RESUME_FOLDER}")

    client = get_client()
    job = extract_job(client, JOB_DESCRIPTION)
    candidates: list[CandidateReport] = []
    failures: list[dict[str, str]] = []

    print(f"Evaluating {len(files)} resume(s) for: {job.role}")
    for file_path in files:
        try:
            print(f"Processing: {file_path.name}")
            text = read_resume(file_path)
            if not text:
                raise ValueError("No readable text was extracted from this file.")
            resume = parse_resume(client, text)
            evaluation = evaluate_candidate(client, job, resume)
            candidates.append(CandidateReport(file_name=file_path.name, candidate=resume, evaluation=evaluation))
        except Exception as error:
            print(f"  Skipped: {error}")
            failures.append({"file_name": file_path.name, "error": str(error)})

    candidates.sort(key=lambda item: item.evaluation.overall_score, reverse=True)
    for rank, candidate in enumerate(candidates, start=1):
        candidate.rank = rank
        print_candidate(candidate)

    report_path = save_report(job, candidates, failures)
    print(f"\nStructured JSON report saved to: {report_path}")


if __name__ == "__main__":
    main()
